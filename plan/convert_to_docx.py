#!/usr/bin/env python3
"""Convert design_report_fixed.txt to a formatted Word .docx file."""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 0, 0)

with open('/home/rupert/Desktop/embedded Projecy/plan/design_report_fixed.txt', 'r') as f:
    text = f.read()

# Split into paragraphs separated by blank lines
blocks = re.split(r'\n\n+', text)

doc.add_heading('IR Remote Control Replacement – Design Report', level=0)

for block in blocks[1:]:  # skip title line
    block = block.rstrip()
    if not block.strip():
        continue

    first_line = block.split('\n')[0]

    # Main section heading: "1. Introduction"
    m = re.match(r'^(\d+)\.\s+(.+)$', first_line)
    if m and '\n' not in block.strip():
        doc.add_heading(f'{m.group(1)}. {m.group(2)}', level=1)
        continue

    # Sub-section heading: "2.1 System Block Diagram"
    m = re.match(r'^(\d+\.\d+)\s+(.+)$', first_line)
    if m and '\n' not in block.strip():
        doc.add_heading(f'{m.group(1)} {m.group(2)}', level=2)
        continue

    # Test heading: "Test 1: ..."
    m = re.match(r'^(Test \d+):\s+(.+)$', first_line)
    if m and '\n' not in block.strip():
        doc.add_heading(f'{m.group(1)}: {m.group(2)}', level=2)
        continue

    # Check if block looks like a diagram/code (most lines are indented 4+)
    blines = block.split('\n')
    indented_count = sum(1 for l in blines if l.startswith('    ') or l.strip() == '')
    has_box_chars = bool(re.search(r'[\+\|=<>\[\]\\\/]{2,}|--[\+\-]--|==', block))
    has_code = bool(re.search(r'void |#define |typedef |while\(|GPTM|Timer0A_Handler|SysTick_Handler|delay_us|uart_|\.valid|\.count|capture_|IRSignal', block))

    if indented_count >= len(blines) * 0.7 and (has_box_chars or has_code):
        # Render as monospaced code/diagram
        for line in blines:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        continue

    # Check if block is a table (has dashed separator line)
    if re.search(r'^  *-{4,}(\s+-{4,})+', block, re.MULTILINE):
        # Find header and separator
        for idx, line in enumerate(blines):
            if re.match(r'^  *-{4,}(\s+-{4,})+', line):
                header_line = blines[idx - 1] if idx > 0 else ''
                sep_line = line
                data_lines = blines[idx + 1:]

                # Find column positions
                cols = []
                in_dash = False
                col_start = 0
                for ci, ch in enumerate(sep_line):
                    if ch == '-' and not in_dash:
                        col_start = ci
                        in_dash = True
                    elif ch != '-' and in_dash:
                        cols.append((col_start, ci))
                        in_dash = False
                if in_dash:
                    cols.append((col_start, len(sep_line)))

                headers = []
                for cs, ce in cols:
                    h = header_line[cs:ce].strip() if cs < len(header_line) else ''
                    headers.append(h)

                rows = []
                for dl in data_lines:
                    if dl.strip() == '':
                        continue
                    row = []
                    for cs, ce in cols:
                        cell = dl[cs:ce].strip() if cs < len(dl) else ''
                        row.append(cell)
                    if any(c for c in row):
                        rows.append(row)

                if headers and rows:
                    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                    table.style = 'Table Grid'
                    for j, h in enumerate(headers):
                        cell = table.rows[0].cells[j]
                        cell.text = h
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.bold = True
                                r.font.size = Pt(10)
                    for ri, row in enumerate(rows):
                        for j, val in enumerate(row):
                            cell = table.rows[ri + 1].cells[j]
                            cell.text = val
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.font.size = Pt(10)
                    doc.add_paragraph()

                # Print any text before the table header as a paragraph
                pre_text = '\n'.join(blines[:max(0, idx - 1)]).strip()
                if pre_text:
                    doc.add_paragraph(pre_text)
                break
        continue

    # Regular text block - join lines and add as paragraph
    # Handle bullet lists within block
    if re.match(r'^  +- ', first_line):
        # Block of bullet points
        items = re.split(r'\n  +- ', block)
        for item in items:
            item = item.lstrip(' -').strip()
            # Join continuation lines
            item = re.sub(r'\n\s+', ' ', item)
            if item:
                doc.add_paragraph(item, style='List Bullet')
        continue

    # Numbered list items
    if re.match(r'^  +\d+\.\s', first_line):
        items = re.split(r'\n  +(?=\d+\.\s)', block)
        for item in items:
            item = item.strip()
            m = re.match(r'\d+\.\s+(.*)', item, re.DOTALL)
            if m:
                txt = re.sub(r'\n\s+', ' ', m.group(1)).strip()
                doc.add_paragraph(txt, style='List Number')
        continue

    # Plain paragraph - collapse whitespace
    text_block = re.sub(r'\n\s*', ' ', block).strip()
    if text_block:
        doc.add_paragraph(text_block)

# Save
output = '/home/rupert/Desktop/embedded Projecy/plan/design_report.docx'
doc.save(output)
print(f'Saved: {output}')
